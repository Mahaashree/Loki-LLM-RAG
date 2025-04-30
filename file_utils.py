import os
import re
import fitz  # PyMuPDF
import docx
import csv
import json
import zipfile
import tempfile
import shutil
from typing import List, Dict, Union, Optional, Tuple
from langchain.schema import Document
from langchain.document_loaders import (
    PyPDFLoader, 
    TextLoader, 
    CSVLoader, 
    UnstructuredWordDocumentLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader
)

def extract_metadata_from_pdf(file_path: str) -> Dict:
    """Extract detailed metadata from PDF files."""
    metadata = {
        "source": file_path,
        "filename": os.path.basename(file_path),
        "filetype": os.path.splitext(file_path)[1],
        "created_at": os.path.getctime(file_path),
        "modified_at": os.path.getmtime(file_path),
        "file_size": os.path.getsize(file_path)
    }
    
    try:
        doc = fitz.open(file_path)
        metadata["page_count"] = len(doc)
        metadata["title"] = doc.metadata.get("title", "")
        metadata["author"] = doc.metadata.get("author", "")
        metadata["subject"] = doc.metadata.get("subject", "")
        metadata["keywords"] = doc.metadata.get("keywords", "")
        
        # Get text statistics
        text = ""
        for page in doc:
            text += page.get_text()
        
        metadata["word_count"] = len(re.findall(r'\w+', text))
        metadata["char_count"] = len(text)
        doc.close()
    except Exception as e:
        print(f"Error extracting PDF metadata: {e}")
    
    return metadata

def extract_metadata_from_docx(file_path: str) -> Dict:
    """Extract detailed metadata from DOCX files."""
    metadata = {
        "source": file_path,
        "filename": os.path.basename(file_path),
        "filetype": os.path.splitext(file_path)[1],
        "created_at": os.path.getctime(file_path),
        "modified_at": os.path.getmtime(file_path),
        "file_size": os.path.getsize(file_path)
    }
    
    try:
        doc = docx.Document(file_path)
        
        # Get core properties
        try:
            core_props = doc.core_properties
            metadata["title"] = core_props.title or ""
            metadata["author"] = core_props.author or ""
            metadata["subject"] = core_props.subject or ""
            metadata["keywords"] = core_props.keywords or ""
        except:
            pass
        
        # Count paragraphs, words, etc.
        text = "\n".join([p.text for p in doc.paragraphs])
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["word_count"] = len(re.findall(r'\w+', text))
        metadata["char_count"] = len(text)
    except Exception as e:
        print(f"Error extracting DOCX metadata: {e}")
    
    return metadata

def extract_metadata_from_csv(file_path: str) -> Dict:
    """Extract detailed metadata from CSV files."""
    metadata = {
        "source": file_path,
        "filename": os.path.basename(file_path),
        "filetype": os.path.splitext(file_path)[1],
        "created_at": os.path.getctime(file_path),
        "modified_at": os.path.getmtime(file_path),
        "file_size": os.path.getsize(file_path)
    }
    
    try:
        with open(file_path, 'r', newline='', encoding='utf-8') as csvfile:
            # Try to determine dialect
            dialect = csv.Sniffer().sniff(csvfile.read(1024))
            csvfile.seek(0)
            
            # Get headers and row count
            reader = csv.reader(csvfile, dialect)
            headers = next(reader, [])
            
            # Count rows
            row_count = sum(1 for _ in reader)
            
            metadata["column_count"] = len(headers)
            metadata["row_count"] = row_count
            metadata["headers"] = headers
    except Exception as e:
        print(f"Error extracting CSV metadata: {e}")
    
    return metadata

def extract_text_from_pdf(file_path: str, include_page_numbers: bool = True) -> List[Document]:
    """Extract text from PDF with enhanced processing."""
    docs = []
    try:
        pdf = fitz.open(file_path)
        
        for i, page in enumerate(pdf):
            # Extract text
            text = page.get_text()
            
            # Create metadata
            metadata = {
                "source": file_path,
                "filename": os.path.basename(file_path),
                "filetype": "pdf",
                "page": i + 1,
                "total_pages": len(pdf)
            }
            
            # Add page number reference if requested
            if include_page_numbers:
                text = f"[Page {i+1}]\n{text}"
            
            # Create document
            doc = Document(page_content=text, metadata=metadata)
            docs.append(doc)
        
        pdf.close()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    
    return docs

def get_file_extension(file_path: str) -> str:
    """Get normalized file extension."""
    return os.path.splitext(file_path)[1].lower()

def is_archive_file(file_path: str) -> bool:
    """Check if file is an archive."""
    extensions = ['.zip', '.tar', '.gz', '.rar', '.7z']
    return get_file_extension(file_path) in extensions

def process_archive_file(file_path: str) -> List[Document]:
    """Extract and process files from archives."""
    docs = []
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Only handling ZIP files for now
        if file_path.endswith('.zip'):
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Process each extracted file
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    extracted_path = os.path.join(root, file)
                    # Skip hidden files and directories
                    if os.path.basename(extracted_path).startswith('.'):
                        continue
                    
                    # Process the file
                    file_docs = process_file(extracted_path)
                    
                    # Add archive source to metadata
                    for doc in file_docs:
                        doc.metadata["archive_source"] = file_path
                    
                    docs.extend(file_docs)
    
    except Exception as e:
        print(f"Error processing archive {file_path}: {e}")
    
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir)
    
    return docs

def process_file(file_path: str) -> List[Document]:
    """Process file based on type with enhanced metadata."""
    ext = get_file_extension(file_path)
    docs = []
    
    # Handle archives
    if is_archive_file(file_path):
        return process_archive_file(file_path)
    
    try:
        # PDF files - use custom extraction
        if ext == '.pdf':
            docs = extract_text_from_pdf(file_path)
        
        # Use specific loaders with enhanced metadata
        elif ext in ['.txt', '.md', '.rst', '.html', '.htm']:
            loader = TextLoader(file_path)
            docs = loader.load()
            # Enhance metadata
            for doc in docs:
                doc.metadata.update({
                    "filetype": ext[1:],
                    "line_count": doc.page_content.count('\n') + 1,
                    "word_count": len(re.findall(r'\w+', doc.page_content))
                })
        
        elif ext in ['.docx', '.doc']:
            loader = UnstructuredWordDocumentLoader(file_path)
            docs = loader.load()
            # Add enhanced metadata
            meta = extract_metadata_from_docx(file_path)
            for doc in docs:
                doc.metadata.update(meta)
        
        elif ext in ['.csv', '.tsv']:
            loader = CSVLoader(file_path)
            docs = loader.load()
            # Add enhanced metadata
            meta = extract_metadata_from_csv(file_path)
            for doc in docs:
                doc.metadata.update(meta)
        
        elif ext in ['.pptx', '.ppt']:
            loader = UnstructuredPowerPointLoader(file_path)
            docs = loader.load()
            # Basic metadata
            for doc in docs:
                doc.metadata.update({
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "filetype": ext[1:]
                })
        
        elif ext in ['.xlsx', '.xls']:
            loader = UnstructuredExcelLoader(file_path)
            docs = loader.load()
            # Basic metadata
            for doc in docs:
                doc.metadata.update({
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "filetype": ext[1:]
                })
        
        else:
            print(f"Unsupported file format: {ext}")
            return []
        
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")
        return []
    
    return docs

def scan_directory(directory_path: str, recursive: bool = True) -> List[str]:
    """Scan directory for processable files."""
    supported_extensions = [
        '.pdf', '.txt', '.docx', '.doc', '.csv', '.pptx', 
        '.ppt', '.xlsx', '.xls', '.md', '.html', '.htm',
        '.zip' # Archive support
    ]
    
    file_paths = []
    
    if recursive:
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                if any(file.lower().endswith(ext) for ext in supported_extensions):
                    file_paths.append(file_path)
    else:
        for file in os.listdir(directory_path):
            file_path = os.path.join(directory_path, file)
            if os.path.isfile(file_path) and any(file.lower().endswith(ext) for ext in supported_extensions):
                file_paths.append(file_path)
    
    return file_paths

def get_document_statistics(file_path: str) -> Dict:
    """Get detailed statistics about a document."""
    ext = get_file_extension(file_path)
    
    if ext == '.pdf':
        return extract_metadata_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_metadata_from_docx(file_path)
    elif ext in ['.csv', '.tsv']:
        return extract_metadata_from_csv(file_path)
    else:
        # Basic stats for other file types
        return {
            "source": file_path,
            "filename": os.path.basename(file_path),
            "filetype": ext[1:] if ext else "unknown",
            "file_size": os.path.getsize(file_path),
            "created_at": os.path.getctime(file_path),
            "modified_at": os.path.getmtime(file_path)
        }

def process_documents_with_metadata(file_paths: List[str]) -> List[Document]:
    """Process multiple documents with enhanced metadata."""
    docs = []
    
    for file_path in file_paths:
        if not os.path.exists(file_path):
            print(f"File does not exist: {file_path}")
            continue
        
        file_docs = process_file(file_path)
        if file_docs:
            print(f"Processed {file_path}: {len(file_docs)} chunks")
            docs.extend(file_docs)
    
    return docs